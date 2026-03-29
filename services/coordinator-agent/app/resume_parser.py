import io
import re
from os import path
from typing import Optional

MAX_RESUME_TEXT_LENGTH = 20000


class ResumeParsingError(Exception):
    def __init__(self, detail: str, *, status_code: int = 422):
        super().__init__(detail)
        self.detail = detail
        self.status_code = status_code


def extract_resume_text(*, filename: Optional[str], content_type: Optional[str], raw: bytes) -> str:
    if not raw:
        raise ResumeParsingError("resume file is empty")

    parser_name = _detect_parser_name(filename=filename, content_type=content_type)
    parser = _PARSERS.get(parser_name)
    if parser is None:
        extension = path.splitext(filename or "")[1].lower() or "unknown"
        raise ResumeParsingError(
            f"unsupported resume file type: {extension}. Supported types: .txt, .pdf, .docx",
            status_code=415,
        )

    try:
        extracted_text = parser(raw)
    except ResumeParsingError:
        raise
    except Exception as exc:
        label = parser_name.upper()
        raise ResumeParsingError(f"failed to extract text from {label} resume: {exc}") from exc

    normalized = _normalize_text(extracted_text)
    if not normalized:
        label = parser_name.upper()
        raise ResumeParsingError(f"{label} resume did not contain any extractable text")
    return normalized[:MAX_RESUME_TEXT_LENGTH]


def _detect_parser_name(*, filename: Optional[str], content_type: Optional[str]) -> Optional[str]:
    normalized_content_type = (content_type or "").split(";", 1)[0].strip().lower()
    if normalized_content_type in _CONTENT_TYPE_TO_PARSER:
        return _CONTENT_TYPE_TO_PARSER[normalized_content_type]

    extension = path.splitext(filename or "")[1].lower()
    return _EXTENSION_TO_PARSER.get(extension)


def _extract_text_from_txt(raw: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ResumeParsingError("failed to decode TXT resume")


def _extract_text_from_pdf(raw: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ResumeParsingError("PDF parsing dependency is not installed") from exc

    reader = PdfReader(io.BytesIO(raw))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_text_from_docx(raw: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ResumeParsingError("DOCX parsing dependency is not installed") from exc

    document = Document(io.BytesIO(raw))
    text_parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)
    return "\n".join(text_parts)


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


_CONTENT_TYPE_TO_PARSER = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
}

_EXTENSION_TO_PARSER = {
    ".docx": "docx",
    ".pdf": "pdf",
    ".txt": "txt",
}

_PARSERS = {
    "docx": _extract_text_from_docx,
    "pdf": _extract_text_from_pdf,
    "txt": _extract_text_from_txt,
}
